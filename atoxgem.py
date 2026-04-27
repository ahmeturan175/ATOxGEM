import os
import re
import time
import shutil
import threading
import requests
import traceback
import sys
import socket
import difflib 
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from docx import Document
from docx.shared import RGBColor
from pypdf import PdfReader
import urllib3

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# =============================================================================
#                           AYARLAR VE RENKLER
# =============================================================================
GOOGLE_BLUE = "#4285F4"
GOOGLE_RED = "#EA4335"
GOOGLE_YELLOW = "#FBBC05"
GOOGLE_GREEN = "#34A853"
BG_COLOR = "#FFFFFF"

stop_event = threading.Event()

# =============================================================================
#                           YARDIMCI FONKSİYONLAR
# =============================================================================

def internet_var_mi():
    try:
        socket.create_connection(("8.8.8.8", 53), timeout=3)
        return True
    except OSError:
        pass
    return False

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def pencereyi_ortala(window, width, height):
    screen_width = window.winfo_screenwidth()
    screen_height = window.winfo_screenheight()
    x = int((screen_width / 2) - (width / 2))
    y = int((screen_height / 2) - (height / 2))
    window.geometry(f"{width}x{height}+{x}+{y}")

def log_yaz(mesaj):
    def _safe_write():
        try:
            text_log.config(state=tk.NORMAL)
            text_log.insert(tk.END, f">> {mesaj}\n")
            text_log.see(tk.END)
            text_log.config(state=tk.DISABLED)
        except:
            pass
    root.after(0, _safe_write)

def metni_sadelestir(metin):
    if not metin: return ""
    metin = metin.lower()
    ceviriler = str.maketrans("ğüşıöçIİâîû", "gusiociiaiu") 
    metin = metin.translate(ceviriler)
    return re.sub(r'[^a-z0-9]', '', metin)

def kelime_temizle(kelime):
    if not kelime: return ""
    ceviriler = str.maketrans("ğüşıöçIİâîû", "gusiociiaiu")
    return re.sub(r'[^a-z0-9]', '', kelime.lower().translate(ceviriler))

def rakamlari_temizle(metin):
    if not metin: return ""
    metin = re.sub(r'\d+', '', metin)
    metin = re.sub(r'[†‡*§]', '', metin)
    return " ".join(metin.split())

def dosya_adindan_bilgi_al(dosya_adi):
    try:
        ad = os.path.splitext(dosya_adi)[0]
        yil_match = re.search(r'\b(19|20)\d{2}\b', ad)
        yil = yil_match.group(0) if yil_match else ""
        temiz_ad = re.sub(r'\b(19|20)\d{2}\b', '', ad)
        temiz_ad = re.sub(r'\b(ve ark\.?|ve arkadaşları|et al\.?|et al)\b', '', temiz_ad, flags=re.IGNORECASE)
        temiz_ad = re.sub(r'[_\-!]', ' ', temiz_ad).strip()
        parcalar = temiz_ad.split()
        ilk_yazar = parcalar[0] if parcalar else ""
        return ilk_yazar, yil
    except:
        return "", ""

def benzersiz_klasor_yolu_bul(temel_yol):
    if not os.path.exists(temel_yol): return temel_yol
    sayac = 1
    while True:
        yeni_yol = f"{temel_yol} ({sayac})"
        if not os.path.exists(yeni_yol): return yeni_yol
        sayac += 1

# =============================================================================
#                           PDF ANALİZ MOTORU
# =============================================================================

def pdf_analiz_et(dosya_yolu, hedef_yazar):
    bulunan_doi = None
    bulunan_metin = None
    try:
        reader = PdfReader(dosya_yolu)
        limit = min(3, len(reader.pages))
        for i in range(limit):
            sayfa = reader.pages[i].extract_text()
            if not bulunan_doi:
                m = re.search(r'\b(10\.\d{4,9}/[-._;()/:a-zA-Z0-9]+)\b', sayfa)
                if m: bulunan_doi = m.group(1)
            
            if not bulunan_metin and hedef_yazar and len(hedef_yazar) > 2:
                if hedef_yazar.lower() in sayfa.lower():
                    satirlar = sayfa.split('\n')
                    for j, s in enumerate(satirlar):
                        if hedef_yazar.lower() in s.lower():
                            start = max(0, j - 2)
                            end = min(len(satirlar), j + 1)
                            parca_listesi = satirlar[start:end]
                            blok = " ".join(parca_listesi)
                            bulunan_metin = rakamlari_temizle(blok)
                            break
            if bulunan_doi and bulunan_metin: break
    except: pass
    return bulunan_doi, bulunan_metin

def derin_tarama_yap(hedef_yazar, hedef_yil, kalan_pdfler, pdf_havuzu):
    if not hedef_yazar or len(hedef_yazar) < 2: return None
    hedef_yazar_lower = hedef_yazar.lower()
    for sade_anahtar in kalan_pdfler:
        pdf_yolu = pdf_havuzu[sade_anahtar]
        try:
            reader = PdfReader(pdf_yolu)
            limit = min(3, len(reader.pages))
            for i in range(limit):
                metin = reader.pages[i].extract_text()
                if not metin: continue
                metin_lower = metin.lower()
                if hedef_yazar_lower in metin_lower:
                    if hedef_yil:
                        if hedef_yil in metin_lower: return sade_anahtar
                    else: return sade_anahtar
        except: pass
    return None

# =============================================================================
#                           API & LİNK TOPLAMA İŞLEMLERİ
# =============================================================================

def crossref_getir(doi):
    try:
        url = f"https://doi.org/{doi}"
        headers = {"Accept": "text/bibliography; style=apa", "User-Agent": "ATÖxGem/1.0"}
        r = requests.get(url, headers=headers, timeout=10)
        if r.status_code == 200: return r.content.decode('utf-8').strip()
    except: pass
    return None

def semantic_ara(sorgu):
    if not sorgu or len(sorgu) < 3: return None
    try:
        url = "https://api.semanticscholar.org/graph/v1/paper/search"
        sorgu_temiz = sorgu[:150].strip()
        params = {"query": sorgu_temiz, "limit": 1, "fields": "title,authors,year,venue"}
        r = requests.get(url, params=params, timeout=5)
        if r.status_code == 429:
            time.sleep(2)
            r = requests.get(url, params=params, timeout=5)
        if r.status_code == 200:
            data = r.json()
            if data.get("total", 0) > 0:
                p = data["data"][0]
                authors = p.get("authors", [])
                yazarlar = [y["name"] for y in authors]
                if not yazarlar: yaz = "Anonim"
                elif len(yazarlar) == 1: yaz = yazarlar[0]
                elif len(yazarlar) == 2: yaz = f"{yazarlar[0]} & {yazarlar[1]}"
                elif len(yazarlar) > 20: yaz = ", ".join(yazarlar[:19]) + ", ... " + yazarlar[-1]
                else: yaz = ", ".join(yazarlar[:-1]) + ", & " + yazarlar[-1]
                baslik = p.get('title', 'Başlıksız')
                dergi = p.get('venue', '')
                yil_v = p.get('year', 'n.d.')
                cikti = f"{yaz} ({yil_v}). {baslik}."
                if dergi: cikti += f" {dergi}."
                return cikti
    except: pass
    return None

def europe_pmc_link_getir(doi):
    try:
        url = f"https://www.ebi.ac.uk/europepmc/webservices/rest/search?query=DOI:{doi}&format=json"
        res = requests.get(url, timeout=7).json()
        results = res.get('resultList', {}).get('result', [])
        if results:
            pmcid = results[0].get('pmcid')
            if pmcid:
                return f"https://europepmc.org/backend/ptpmcrender.fcgi?accid={pmcid}&blobtype=pdf"
    except: pass
    return None

def link_topla(sorgu_metni, custom_url):
    if not sorgu_metni or len(sorgu_metni) < 5: 
        return "Bilinmeyen Metin", None, 0
    try:
        doi_match = re.search(r'\b(10\.\d{4,9}/[-._;()/:a-zA-Z0-9]+)\b', sorgu_metni)
        if doi_match:
            doi = doi_match.group(1)
            aranan_isim = f"DOI: {doi}"
            if custom_url:
                c_url = custom_url.strip()
                if c_url:
                    if not c_url.endswith("/"): c_url += "/"
                    return aranan_isim, f"{c_url}{doi}", 3 
            return aranan_isim, f"https://doi.org/{doi}", 3

        baslik_match = re.search(r'\(\d{4}[a-z]?\)\.?\s*(.+?)(?:\.|$)', sorgu_metni)
        if baslik_match:
            sorgu_temiz = baslik_match.group(1).strip()
        else:
            sorgu_temiz = re.sub(r'[^\w\s]', ' ', sorgu_metni)
        
        sorgu_temiz = " ".join(sorgu_temiz.split()[:15])
        aranan_isim = sorgu_temiz 
        
        url_search = "https://api.semanticscholar.org/graph/v1/paper/search"
        params = {"query": sorgu_temiz, "limit": 1, "fields": "title,openAccessPdf,externalIds"}
        r = requests.get(url_search, params=params, timeout=7)
        if r.status_code == 429: 
            time.sleep(2)
            r = requests.get(url_search, params=params, timeout=7)
            
        if r.status_code == 200:
            data = r.json()
            if data.get("total", 0) > 0:
                paper = data["data"][0]
                ext_ids = paper.get("externalIds", {})
                if "DOI" in ext_ids:
                    doi = ext_ids["DOI"]
                    if custom_url:
                        c_url = custom_url.strip()
                        if c_url:
                            if not c_url.endswith("/"): c_url += "/"
                            return aranan_isim, f"{c_url}{doi}", 2
                    
                    pmc_link = europe_pmc_link_getir(doi)
                    if pmc_link: return aranan_isim, pmc_link, 2
                    
                    unpaywall_url = f"https://api.unpaywall.org/v2/{doi}?email=atoxgem_research@example.com"
                    try:
                        u_req = requests.get(unpaywall_url, timeout=7)
                        if u_req.status_code == 200:
                            u_data = u_req.json()
                            if u_data.get("is_oa"):
                                locations = u_data.get("oa_locations", [])
                                for loc in locations:
                                    url_pdf = loc.get("url_for_pdf")
                                    if url_pdf: return aranan_isim, url_pdf, 2
                    except: pass
                
                oa_info = paper.get("openAccessPdf")
                if oa_info and oa_info.get("url"):
                    return aranan_isim, oa_info["url"], 2
        return aranan_isim, None, 0
    except: 
        fallback_isim = " ".join(sorgu_metni.split()[:5]) + "..."
        return fallback_isim, None, 0

# =============================================================================
#                           ANA İŞLEM ZİNCİRİ
# =============================================================================

def tam_otomatik_islem(word_yolu, kaynak_pdf_klasoru, indirme_aktif, apa_aktif, custom_url, kontrol_aktif):
    try:
        stop_event.clear()
        if stop_event.is_set(): return
        log_yaz("1. Hazırlık yapılıyor... / Preparing...")
        
        is_standalone_apa = (not kontrol_aktif) and apa_aktif and kaynak_pdf_klasoru and not word_yolu and not indirme_aktif
        
        if is_standalone_apa:
            is_adi = os.path.basename(kaynak_pdf_klasoru.rstrip('/\\')) + "_Sadece_APA"
            word_tam_isim = is_adi
        else:
            is_adi = os.path.splitext(os.path.basename(word_yolu))[0].strip()
            word_tam_isim = os.path.basename(word_yolu)
            
        is_adi = re.sub(r'[<>:"/\\|?*]', '_', is_adi)
        masaustu = os.path.join(os.path.expanduser("~"), "Desktop")
        
        temel_klasor_adi = f"ATÖxGem_{is_adi}"
        temel_yol = os.path.join(masaustu, temel_klasor_adi)
        ana_cikis_yolu = benzersiz_klasor_yolu_bul(temel_yol)
        ana_klasor_adi = os.path.basename(ana_cikis_yolu)
        os.makedirs(ana_cikis_yolu)
        
        def _apa_uret(hedef_dizin):
            doc_ref = Document()
            doc_ref.add_heading('Kaynakça (APA 7)', 0)
            hedef_dosyalar = [f for f in os.listdir(hedef_dizin) if f.lower().endswith('.pdf')]
            for i, dosya in enumerate(hedef_dosyalar):
                if stop_event.is_set(): break
                log_yaz(f"  > APA: {dosya}")
                tam_yol = os.path.join(hedef_dizin, dosya)
                bilgi = dosya_adindan_bilgi_al(dosya)
                yazar_ipucu = bilgi[0]
                yil_ipucu = bilgi[1]
                apa_metni = None
                renk = RGBColor(255, 0, 0)
                
                analiz_sonuc = pdf_analiz_et(tam_yol, yazar_ipucu)
                doi = analiz_sonuc[0]
                metin = analiz_sonuc[1]
                
                if doi:
                    apa_metni = crossref_getir(doi)
                    if apa_metni: renk = RGBColor(0, 128, 0) 
                if not apa_metni:
                    if metin:
                        apa_metni = semantic_ara(metin)
                    if not apa_metni and yazar_ipucu:
                        sorgu_isim = f"{yazar_ipucu} {yil_ipucu}".strip()
                        apa_metni = semantic_ara(sorgu_isim)
                    if apa_metni: renk = RGBColor(204, 153, 0) 
                if not apa_metni: 
                    apa_metni = "BULUNAMADI / NOT FOUND"
                    renk = RGBColor(255, 0, 0) 
                
                p = doc_ref.add_paragraph()
                run_dosya = p.add_run(f"[{dosya}]: ")
                run_dosya.bold = True
                run_dosya.font.color.rgb = RGBColor(0, 0, 0)
                run_apa = p.add_run(apa_metni)
                run_apa.font.color.rgb = renk
                if renk == RGBColor(255, 0, 0): run_apa.bold = True
                time.sleep(0.3)

            if not stop_event.is_set():
                doc_ref.save(os.path.join(ana_cikis_yolu, "3_APA_Kaynakca.docx"))

        if is_standalone_apa:
            _apa_uret(kaynak_pdf_klasoru)
        else:
            doc = Document(word_yolu)
            if kontrol_aktif:
                log_yaz("-> Çapraz Kontrol / Cross Check...")
                kaynakca_basladi = False
                kaynakca_listesi = []
                icerik_paragraflari = []
                
                for p in doc.paragraphs:
                    text = p.text.strip()
                    if not text: continue
                    if not kaynakca_basladi:
                        if re.match(r'^\s*(KAYNAKÇA|KAYNAKLAR|REFERENCES)\b', text, re.IGNORECASE):
                            kaynakca_basladi = True
                            continue
                        icerik_paragraflari.append(p)
                    else:
                        if len(text) > 5:
                            kaynakca_listesi.append({"p": p, "orijinal": text, "lower": text.lower(), "durum": None, "ref_id": None})
                
                # YENİ: Otomatik Format Algılama (APA vs Numaralı)
                numbered_count = 0
                for kaynak in kaynakca_listesi:
                    # '20.', '[21]', '1)' gibi formatları algılar
                    m = re.match(r'^\s*\[?(\d+)\]?[\.\-\)]?\s+', kaynak["orijinal"])
                    if m:
                        kaynak["ref_id"] = int(m.group(1))
                        numbered_count += 1

                is_numbered_style = False
                if kaynakca_listesi and (numbered_count / len(kaynakca_listesi)) > 0.4:
                    is_numbered_style = True
                    log_yaz("-> Numaralı Atıf Sistemi Tespit Edildi (Vancouver/IEEE)")
                else:
                    log_yaz("-> Yazar-Yıl Atıf Sistemi Tespit Edildi (APA vb.)")

                hatali_atiflar = set()
                kismi_hatali_atiflar = set()
                
                if is_numbered_style:
                    # [1, 2] veya [33-35] formatları için
                    re_parantez_num = r'(\[\s*\d+(?:\s*[,;\-]\s*\d+)*\s*\])'
                    for paragraf in icerik_paragraflari:
                        if stop_event.is_set(): break
                        ham = paragraf.text
                        if "[" not in ham: continue
                        ornek_run = paragraf.runs[0] if paragraf.runs else None
                        parcalar = re.split(re_parantez_num, ham)
                        if len(parcalar) == 1: continue
                        
                        paragraf.clear()
                        for parca in parcalar:
                            if re.match(re_parantez_num, parca):
                                icerik = parca.strip('[] \t')
                                numaralar = set()
                                for p_num in re.split(r'[,;]', icerik):
                                    if '-' in p_num:
                                        try:
                                            parts = p_num.split('-')
                                            if len(parts) == 2:
                                                numaralar.update(range(int(parts[0].strip()), int(parts[1].strip())+1))
                                        except: pass
                                    else:
                                        try: numaralar.add(int(p_num.strip()))
                                        except: pass
                                
                                if not numaralar:
                                    r = paragraf.add_run(parca)
                                    if ornek_run: r.font.name = ornek_run.font.name
                                    continue
                                
                                bulunanlar = []
                                bulunamayanlar = []
                                
                                for num in numaralar:
                                    found = False
                                    for k in kaynakca_listesi:
                                        if k["ref_id"] == num:
                                            k["durum"] = "tam"
                                            found = True
                                            break
                                    if found: bulunanlar.append(num)
                                    else: bulunamayanlar.append(num)
                                
                                if not bulunamayanlar:
                                    renk = 4 # Hepsi var (Yeşil)
                                elif bulunamayanlar and bulunanlar:
                                    renk = 7 # Kısmi var (Sarı)
                                    kismi_hatali_atiflar.add(f"{parca} -> Bulunamayanlar: {bulunamayanlar}")
                                else:
                                    renk = 6 # Hiçbiri yok (Kırmızı)
                                    hatali_atiflar.add(parca)
                                
                                rr = paragraf.add_run(parca)
                                if ornek_run: rr.font.name = ornek_run.font.name
                                rr.font.highlight_color = renk
                                rr.bold = True
                            else:
                                r = paragraf.add_run(parca)
                                if ornek_run: r.font.name = ornek_run.font.name
                
                else:
                    # Klasik APA (Yazar, Yıl) Modu
                    re_parantez = r'(\(.*?\))'
                    re_yil = r'\b(?:19|20)\d{2}[a-z]?\b'
                    for paragraf in icerik_paragraflari:
                        if stop_event.is_set(): break
                        ham = paragraf.text
                        if "(" not in ham: continue
                        ornek_run = paragraf.runs[0] if paragraf.runs else None
                        parcalar = re.split(re_parantez, ham)
                        paragraf.clear()
                        
                        for parca in parcalar:
                            if parca.startswith("(") and parca.endswith(")"):
                                icerik = parca[1:-1]
                                yil_iter = re.finditer(re_yil, icerik)
                                yeni_icerik = []
                                son_idx = 0
                                for match in yil_iter:
                                    bitis = match.end()
                                    alt = icerik[son_idx:bitis]
                                    son_idx = bitis
                                    renk = None 
                                    if re.search(r'[a-zA-Z]', alt):
                                        temiz_ad = alt.strip(" ;.,")
                                        yil_m = re.search(r'\b(19|20)\d{2}\b', alt)
                                        h_yil = yil_m.group(0) if yil_m else ""
                                        sadece_harf = re.sub(r'[^a-zA-ZğüşıöçĞÜŞİÖÇ\s]', ' ', alt)
                                        kelimeler = sadece_harf.split()
                                        h_yazar = kelimeler[0].lower() if kelimeler else ""
                                        
                                        if h_yazar and h_yil:
                                            h_yazar_temiz = kelime_temizle(h_yazar)
                                            h_yil_int = int(h_yil) if h_yil.isdigit() else 0
                                            eslesti = False
                                            kismi_eslesti = False
                                            
                                            for kaynak in kaynakca_listesi:
                                                if h_yazar in kaynak["lower"] and h_yil in kaynak["lower"]:
                                                    eslesti = True
                                                    kaynak["durum"] = "tam" 
                                                    break
                                            
                                            if not eslesti:
                                                for kaynak in kaynakca_listesi:
                                                    k_metin = kaynak["lower"]
                                                    k_yillar = [int(y) for y in re.findall(r'\b(?:19|20)\d{2}\b', k_metin)]
                                                    k_kelimeler = [kelime_temizle(w) for w in re.findall(r'[a-zA-ZğüşıöçĞÜŞİÖÇ]+', k_metin)[:10]]
                                                    yazar_ok = False
                                                    for k_kelime in k_kelimeler:
                                                        if k_kelime and h_yazar_temiz:
                                                            if difflib.SequenceMatcher(None, h_yazar_temiz, k_kelime).ratio() >= 0.80:
                                                                yazar_ok = True
                                                                break
                                                    yil_ok = False
                                                    yil_yaklasik = False
                                                    if h_yil_int in k_yillar: yil_ok = True
                                                    else:
                                                        for ky in k_yillar:
                                                            if abs(ky - h_yil_int) <= 2: 
                                                                yil_yaklasik = True
                                                                break
                                                    if yazar_ok and yil_ok:
                                                        eslesti = True
                                                        kaynak["durum"] = "tam"
                                                        break
                                                    elif yazar_ok and yil_yaklasik:
                                                        kismi_eslesti = True
                                                        if kaynak["durum"] != "tam": kaynak["durum"] = "kismi"

                                            if eslesti: renk = 4 
                                            elif kismi_eslesti:
                                                renk = 7 
                                                kismi_hatali_atiflar.add(temiz_ad)
                                            else:
                                                renk = 6 
                                                hatali_atiflar.add(temiz_ad)
                                    yeni_icerik.append((alt, renk))
                                if son_idx < len(icerik): yeni_icerik.append((icerik[son_idx:], None))
                                r = paragraf.add_run("(")
                                if ornek_run: r.font.name = ornek_run.font.name
                                for item_tuple in yeni_icerik:
                                    txt = item_tuple[0]
                                    clr = item_tuple[1]
                                    rr = paragraf.add_run(txt)
                                    if ornek_run: rr.font.name = ornek_run.font.name
                                    if clr: 
                                        rr.font.highlight_color = clr 
                                        rr.bold = True
                                r = paragraf.add_run(")")
                                if ornek_run: r.font.name = ornek_run.font.name
                            else:
                                r = paragraf.add_run(parca)
                                if ornek_run: r.font.name = ornek_run.font.name

                kullanilmayan_kaynaklar = []
                for kaynak in kaynakca_listesi:
                    if kaynak["durum"] is None:
                        # Numaralı sistemse hata raporuna o numarayı daha net yazdıralım
                        if is_numbered_style and kaynak["ref_id"] is not None:
                            kullanilmayan_kaynaklar.append(f"[{kaynak['ref_id']}] {kaynak['orijinal'][:40]}...")
                        else:
                            kullanilmayan_kaynaklar.append(kaynak["orijinal"])
                        for run in kaynak["p"].runs:
                            if run.text.strip(): run.font.highlight_color = 6 
                    elif kaynak["durum"] == "kismi":
                        for run in kaynak["p"].runs:
                            if run.text.strip(): run.font.highlight_color = 7 

                if hatali_atiflar or kullanilmayan_kaynaklar or kismi_hatali_atiflar:
                    rapor_yolu = os.path.join(ana_cikis_yolu, "Capraz_Kontrol_Hata_Raporu.txt")
                    with open(rapor_yolu, "w", encoding="utf-8") as f:
                        f.write("ÇAPRAZ KONTROL HATA RAPORU / CROSS-CHECK ERROR REPORT\n")
                        f.write("="*60 + "\n\n")
                        f.write(f"1. METİNDE OLUP KAYNAKÇADA OLMAYANLAR / NOT IN REFERENCES (KIRMIZI/RED - {len(hatali_atiflar)})\n")
                        f.write("-" * 60 + "\n")
                        if hatali_atiflar:
                            for h in sorted(list(hatali_atiflar)): f.write(f"- {h}\n")
                        else: f.write("- Kusursuz / Perfect!\n")
                        f.write("\n\n")
                        f.write(f"2. KISMİ EŞLEŞMELER / PARTIAL MATCHES (SARI/YELLOW - {len(kismi_hatali_atiflar)})\n")
                        f.write("-" * 60 + "\n")
                        if kismi_hatali_atiflar:
                            for h in sorted(list(kismi_hatali_atiflar)): f.write(f"- {h} (Eksik veya Hatalı / Error detected)\n")
                        else: f.write("- Sorun yok / No issues!\n")
                        f.write("\n\n")
                        f.write(f"3. KAYNAKÇADA OLUP KULLANILMAYANLAR / UNUSED REFERENCES (KIRMIZI/RED - {len(kullanilmayan_kaynaklar)})\n")
                        f.write("-" * 60 + "\n")
                        if kullanilmayan_kaynaklar:
                            for k in kullanilmayan_kaynaklar: f.write(f"- {k}\n")
                        else: f.write("- Kusursuz / Perfect!\n")
                
                kayit_yolu = os.path.join(ana_cikis_yolu, "1_Cift_Yonlu_Kontrol_Edilmis_Metin.docx")
                doc.save(kayit_yolu)

            else:
                bulunan_pdf_klasoru = ""
                if kaynak_pdf_klasoru:
                    bulunan_pdf_klasoru = os.path.join(ana_cikis_yolu, "1_Bulunan_PDFler")
                    os.makedirs(bulunan_pdf_klasoru)
                    log_yaz("-> Arşiv Taranıyor / Scanning Archive...")
                    pdf_havuzu = {} 
                    pdf_orijinal = {}
                    walk_generator = os.walk(kaynak_pdf_klasoru)
                    for item in walk_generator:
                        if stop_event.is_set(): break
                        kok = item[0]
                        try: dosyalar = item[2]
                        except: continue
                        for d in dosyalar:
                            if d.lower().endswith('.pdf'):
                                sade = metni_sadelestir(os.path.splitext(d)[0])
                                pdf_havuzu[sade] = os.path.join(kok, d)
                                pdf_orijinal[sade] = os.path.splitext(d)[0]
                    
                    kopyalananlar = set()
                    kayiplar = {} 
                    re_parantez = r'(\(.*?\))'
                    re_yil = r'\b(?:19|20)\d{2}[a-z]?\b'
                    for paragraf in doc.paragraphs:
                        if stop_event.is_set(): break
                        ham = paragraf.text
                        if "(" not in ham: continue
                        ornek_run = paragraf.runs[0] if paragraf.runs else None
                        parcalar = re.split(re_parantez, ham)
                        paragraf.clear()
                        for parca in parcalar:
                            if parca.startswith("(") and parca.endswith(")"):
                                icerik = parca[1:-1]
                                yil_iter = re.finditer(re_yil, icerik)
                                yeni_icerik = []
                                son_idx = 0
                                for match in yil_iter:
                                    bitis = match.end()
                                    alt = icerik[son_idx:bitis]
                                    son_idx = bitis
                                    renk = None
                                    if re.search(r'[a-zA-Z]', alt):
                                        sade = metni_sadelestir(alt)
                                        if len(sade) > 2:
                                            if sade in pdf_havuzu:
                                                if sade not in kopyalananlar:
                                                    hedef = os.path.join(bulunan_pdf_klasoru, pdf_orijinal[sade] + ".pdf")
                                                    shutil.copy2(pdf_havuzu[sade], hedef)
                                                    kopyalananlar.add(sade)
                                                renk = 7 if metni_sadelestir(pdf_orijinal[sade]) == sade else 6
                                            else:
                                                renk = 2
                                                temiz_ad = alt.strip(" ;.,")
                                                yil_m = re.search(r'\b(19|20)\d{2}\b', alt)
                                                h_yil = yil_m.group(0) if yil_m else ""
                                                sadece_harf = re.sub(r'[^a-zA-ZğüşıöçĞÜŞİÖÇ\s]', ' ', alt)
                                                kelimeler = sadece_harf.split()
                                                h_yazar = kelimeler[0] if kelimeler else ""
                                                kayiplar[temiz_ad] = {"yazar": h_yazar, "yil": h_yil}
                                    yeni_icerik.append((alt, renk))
                                if son_idx < len(icerik): yeni_icerik.append((icerik[son_idx:], None))
                                r = paragraf.add_run("(")
                                if ornek_run: r.font.name = ornek_run.font.name
                                for item_tuple in yeni_icerik:
                                    txt = item_tuple[0]
                                    clr = item_tuple[1]
                                    rr = paragraf.add_run(txt)
                                    if ornek_run: rr.font.name = ornek_run.font.name
                                    if clr == 7: rr.font.highlight_color = 4 
                                    elif clr == 6: rr.font.highlight_color = 7 
                                    elif clr == 2: rr.font.highlight_color = 6 
                                    if clr: rr.bold = True
                                r = paragraf.add_run(")")
                                if ornek_run: r.font.name = ornek_run.font.name
                            else:
                                r = paragraf.add_run(parca)
                                if ornek_run: r.font.name = ornek_run.font.name

                    sonradan_bulunanlar = {}
                    if kayiplar:
                        kalan_pdfler = set(pdf_havuzu.keys()) - kopyalananlar
                        for kayip_metin, veriler in kayiplar.items():
                            if stop_event.is_set(): break
                            bulunan_anahtar = derin_tarama_yap(veriler["yazar"], veriler["yil"], kalan_pdfler, pdf_havuzu)
                            if bulunan_anahtar:
                                eski_ad = pdf_orijinal[bulunan_anahtar]
                                yeni_ad = f"!{eski_ad}" 
                                hedef_yol = os.path.join(bulunan_pdf_klasoru, yeni_ad + ".pdf")
                                shutil.copy2(pdf_havuzu[bulunan_anahtar], hedef_yol)
                                kopyalananlar.add(bulunan_anahtar)
                                kalan_pdfler.remove(bulunan_anahtar) 
                                sonradan_bulunanlar[kayip_metin] = yeni_ad

                    if kayiplar:
                        with open(os.path.join(ana_cikis_yolu, "2_Arsivde_Bulunamayan_Atiflar.txt"), "w", encoding="utf-8") as f:
                            f.write("BULUNAMAYAN KAYNAKLAR / MISSING SOURCES:\n")
                            f.write("------------------------------------------\n")
                            for k in sorted(list(kayiplar.keys())):
                                if k in sonradan_bulunanlar: f.write(f"- {k} (İçi Taramada Bulundu: {sonradan_bulunanlar[k]}.pdf)\n")
                                else: f.write(f"- {k} (BULUNAMADI / NOT FOUND)\n")

                if apa_aktif and bulunan_pdf_klasoru and os.path.exists(bulunan_pdf_klasoru):
                    _apa_uret(bulunan_pdf_klasoru)

                if indirme_aktif:
                    log_yaz("-> Linkler Toplanıyor / Collecting Links...")
                    kaynakca_basladi = False
                    toplam_kaynak_sayisi = 0
                    toplanan_linkler = []
                    bulunamayan_linkler = [] 
                    
                    for p in doc.paragraphs:
                        if stop_event.is_set(): break
                        text = p.text.strip()
                        if not text: continue
                        if not kaynakca_basladi:
                            if re.match(r'^\s*(KAYNAKÇA|KAYNAKLAR|REFERENCES)\b', text, re.IGNORECASE):
                                kaynakca_basladi = True
                            continue
                        
                        if kaynakca_basladi and len(text) > 15:
                            toplam_kaynak_sayisi += 1
                            yazar, yil = dosya_adindan_bilgi_al(text[:50])
                            dosya_adi = f"{yazar}_{yil}" if yazar else f"Makale_{int(time.time())}"
                            dosya_adi = re.sub(r'[^a-zA-Z0-9_]', '', dosya_adi)
                            aranan_metin, bulunan_link, yontem = link_topla(text, custom_url)
                            if bulunan_link:
                                toplanan_linkler.append((dosya_adi, aranan_metin, bulunan_link, yontem))
                                renk_kodu = 4 if yontem == 3 else 7 
                            else:
                                bulunamayan_linkler.append(text) 
                                renk_kodu = 6 
                            for run in p.runs:
                                if run.text.strip(): run.font.highlight_color = renk_kodu
                    
                    if bulunamayan_linkler:
                        txt_yolu = os.path.join(ana_cikis_yolu, "4_İnternette_Bulunamayan_Makaleler.txt")
                        with open(txt_yolu, "w", encoding="utf-8") as f:
                            f.write(f"'{word_tam_isim}' BULUNAMAYANLAR / NOT FOUND IN INTERNET\n")
                            f.write("=" * 60 + "\n\n")
                            for eksik in bulunamayan_linkler: f.write(f"- {eksik}\n\n")

                    if toplanan_linkler:
                        html_yolu = os.path.join(ana_cikis_yolu, "5_Toplu_Indirme_Paneli.html")
                        js_link_array = ",\n".join([f'"{link}"' for _, _, link, _ in toplanan_linkler])
                        html_icerik = f"""
                        <html>
                        <head>
                            <meta charset="utf-8">
                            <title>ATÖxGem Toplu İndirme Paneli</title>
                            <style>
                                body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 40px; background-color: #f1f3f4; color: #202124; }}
                                .header-box {{ background: white; padding: 25px; border-radius: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); text-align: center; margin-bottom: 30px; }}
                                h1 {{ margin: 0; color: #4285F4; font-size: 28px; }}
                                .istatistik-text {{ font-size: 18px; color: #EA4335; font-weight: bold; margin-top: 15px; background: #fce8e6; padding: 10px; border-radius: 6px; display: inline-block; }}
                                .master-btn {{ display: block; width: 100%; max-width: 400px; margin: 20px auto; padding: 18px 20px; background-color: #34A853; color: white; border: none; border-radius: 8px; font-size: 18px; font-weight: bold; cursor: pointer; }}
                                .master-btn:hover {{ background-color: #2b8c45; }}
                                .warning {{ font-size: 14px; color: #EA4335; font-weight: bold; margin-bottom: 10px; }}
                                .link-card {{ background: white; padding: 15px; margin-bottom: 10px; border-radius: 8px; border-left: 5px solid #FBBC05; display: flex; justify-content: space-between; align-items: center; }}
                                .link-card.doi-dogrudan {{ border-left: 5px solid #34A853; }}
                                .info-area {{ flex-grow: 1; }}
                                .dosya-adi {{ font-weight: bold; font-size: 16px; color: #202124; }}
                                .aranan {{ font-size: 13px; color: #5f6368; margin-top: 4px; }}
                                .doi-uyari {{ font-size: 12px; color: #EA4335; margin-top: 6px; font-style: italic; font-weight: bold; }}
                                .indirme-butonu {{ padding: 8px 15px; background-color: #4285f4; color: white; text-decoration: none; border-radius: 4px; font-weight: bold; font-size: 14px; }}
                            </style>
                            <script>
                                function tumunuAc() {{
                                    var linkler = [{js_link_array}];
                                    for (var i = 0; i < linkler.length; i++) {{ window.open(linkler[i], '_blank'); }}
                                    alert("Sekmeler açılıyor! Lütfen Pop-up'lara izin verin. / Allow pop-ups!");
                                }}
                            </script>
                        </head>
                        <body>
                            <div class="header-box">
                                <h1>🚀 Toplu İndirme Kontrol Paneli</h1>
                                <button class="master-btn" onclick="tumunuAc()">TÜM PDF'LERİ TEK TIKLA AÇ</button>
                                <p class="warning">⚠️ DİKKAT: Tarayıcınız pop-up'ları engelleyebilir. Adres çubuğundan izin verin.</p>
                                <div class="istatistik-text">
                                    '{word_tam_isim}' dosyasındaki {toplam_kaynak_sayisi} kaynaktan {len(toplanan_linkler)} kadarı internette bulundu!
                                </div>
                            </div>
                            <h3>Bulunan Linkler</h3>
                        """
                        for dosya, aranan, link, yontem in toplanan_linkler:
                            uyari_html = ""
                            css_class = "link-card"
                            if yontem == 3:
                                css_class += " doi-dogrudan"
                                uyari_html = '<div class="doi-uyari">⚠️ Uyarı: DOI numarası hatalıysa, yanlış bir makaleye gidiyor olabilirsiniz.</div>'
                            html_icerik += f"""
                            <div class="{css_class}">
                                <div class="info-area">
                                    <div class="dosya-adi">{dosya}</div>
                                    <div class="aranan">{aranan}</div>
                                    {uyari_html}
                                </div>
                                <a class="indirme-butonu" href="{link}" target="_blank">Aç</a>
                            </div>
                            """
                        html_icerik += "</body></html>"
                        with open(html_yolu, "w", encoding="utf-8") as f:
                            f.write(html_icerik)

                if stop_event.is_set(): return
                kayit_yolu = os.path.join(ana_cikis_yolu, "1_Gorsel_Atif_ve_Kaynakca_Analizi.docx")
                doc.save(kayit_yolu)

        log_yaz("İŞLEM BİTTİ. / DONE.")
        
        s_msg = TRANSLATIONS[lang_var.get().split(" ")[1]]["msg_success"]
        root.after(0, lambda: messagebox.showinfo("ATÖxGem", s_msg))
        root.after(0, lambda: os.startfile(ana_cikis_yolu))
        
    except Exception as e:
        hata = traceback.format_exc()
        rapor_yolu = os.path.join(os.path.expanduser("~"), "Desktop", "HATA_RAPORU.txt")
        with open(rapor_yolu, "w", encoding="utf-8") as f: f.write(hata)
        log_yaz(f"KRİTİK HATA: {e}")
        root.after(0, lambda: messagebox.showerror("Hata/Error", f"Program bir hata ile karşılaştı.\nDetaylar masaüstünde 'HATA_RAPORU.txt' dosyasında."))
    finally:
        root.after(0, reset_ui)

# =============================================================================
#                           DİL VE ÇEVİRİ MOTORU (GUI İÇİN)
# =============================================================================

TRANSLATIONS = {
    "TR": {
        "word_lbl": "1. Word Dosyası:",
        "pdf_lbl": "2. PDF Arşiv Klasörü:",
        "btn_select": "Seç",
        "chk_apa": "Arşivden Bulunan PDF'lerin APA Kaynakçasını Çıkar",
        "chk_indirme": "KAYNAKÇA listesini tara ve İndirme Linklerini Topla",
        "lbl_custom": "Özel Kaynak Linki (Örn: https://.../):",
        "chk_kontrol": "Sadece Atıf / Kaynakça Çapraz Kontrolü Yap",
        "lbl_kontrol_desc": "(Bu mod seçildiğinde üstteki işlemler iptal olur, metin ve kaynakça eşleştirilir)",
        "btn_start": "BAŞLAT",
        "btn_cancel": "İPTAL ET",
        "msg_missing_word": "Lütfen bir Word dosyası seçiniz.\n(Sadece APA çıkaracaksanız Word seçmeden PDF Klasörünü seçebilirsiniz)",
        "msg_missing_word_cross": "Çapraz Kontrol için lütfen bir Word dosyası seçiniz.",
        "msg_missing_pdf": "Lütfen ya PDF klasörü seçin ya da Link Toplama'yı işaretleyin.",
        "msg_no_internet": "Seçtiğiniz işlemler (APA Çıkarma / Link Toplama) için internet bağlantısı gereklidir.",
        "msg_cancel_q": "İşlemi durdurmak istiyor musunuz?",
        "msg_success": "İşlem Tamamlandı!",
        "help_title": "ATÖxGem Kullanım Kılavuzu",
        "help_text": "ATÖxGem v7.0 (Absolut) 🎓\n\nBu program, tez ve makale yazım sürecinde saatler sürecek angarya işleri saniyeler içinde sizin için yapar.\n\n==================================================\nMOD 1: ÇAPRAZ KONTROL MODU (Kırmızı Kutu Seçiliyse)\n==================================================\nAmacı: Metin içindeki atıfların (APA veya Numaralı [1, 2] formatında) gerçekten KAYNAKÇA listesinde olup olmadığını bulmaktır. Program stili otomatik algılar! (Çevrimdışı çalışabilir).\n\nÇıkan Dosyalar Nelerdir?\n1_Cift_Yonlu_Kontrol_Edilmis_Metin.docx:\n   🟩 Yeşil: Atıf, kaynakçada eksiksiz bulundu.\n   🟨 Sarı: Atıfta ufak bir harf hatası, numarada eksiklik veya 1-2 yıllık sapma var.\n   🟥 Kırmızı: Atıf yapılıp kaynağı unutulanlar veya listeye eklenip metinde kullanılmayan Hayalet Kaynaklar.\n\nCapraz_Kontrol_Hata_Raporu.txt:\n   Hataların detaylı listesi.\n\n==================================================\nMOD 2: ARŞİV TARAMA VE İNTERNETTEN LİNK TOPLAMA \n==================================================\nÖzellik A) PDF Arşivi Eşleştirme (Word + Klasör Seçilirse):\nMetindeki atıfları okur ve arşivinizde arar. Bulduğu PDF'leri \"1_Bulunan_PDFler\" klasörüne kopyalar. Ayrıca eşleştirmeyi kontrol edebilmeniz için renkli bir Word dosyası (1_Gorsel_Atif_ve_Kaynakca_Analizi.docx) üretir:\n   🟩 Yeşil (Metin İçi): Atıfın PDF'i arşivde bulundu.\n   🟥 Kırmızı (Metin İçi): Atıfın PDF'i arşivde BULUNAMADI.\n\nÖzellik B) APA Kaynakça Çıkarıcı:\nKlasörden çekilmiş olan PDF'lerin içeriğini internetten doğrulayarak APA 7 formatında Word dosyası hazırlar. (İnternet gerektirir).\nRenk Kodları (APA Word Dosyası İçin):\n   🟩 Yeşil: Kaynak PDF'in içindeki DOI kodundan okundu ve %100 eşleşti.\n   🟨 Sarı: Kaynak makalenin/yazarın adından aratılarak bulundu.\n   🟥 Kırmızı: Kaynak internette BULUNAMADI.\n\nÖzellik C) İndirme Linklerini Toplama:\nKAYNAKÇA listenizi okur ve her bir makalenin indirme linkini internette arar. Tarayıcıda şelale gibi indiren HTML paneli hazırlar.\nRenk Kodları (Ana Word Dosyası Kaynakçası İçin):\n   🟩 Yeşil: Kaynak internette \"DOI Numarası\" ile kesin olarak bulundu.\n   🟨 Sarı: DOI yoktu, \"Başlığından\" aratarak buldu.\n   🟥 Kırmızı: Makale internette bulunamadı.\n\n==================================================\nMOD 3: BAĞIMSIZ APA MODU \n==================================================\nHiç Word dosyası seçmeden sadece \"PDF Arşiv Klasörü\" seçip \"APA Kaynakçasını Çıkar\" işaretlerseniz, o klasördeki tüm PDF'lerin APA kaynakçasını çıkarır.\n\nÇalışmalarınızda başarılar dileriz. 🚀"
    },
    "EN": {
        "word_lbl": "1. Word File:",
        "pdf_lbl": "2. PDF Archive Folder:",
        "btn_select": "Select",
        "chk_apa": "Extract APA References from Found PDFs",
        "chk_indirme": "Scan REFERENCE list and Collect Download Links",
        "lbl_custom": "Custom Source Link (e.g. https://.../):",
        "chk_kontrol": "Only Perform Citation / Reference Cross-Check",
        "lbl_kontrol_desc": "(When this mode is selected, the operations above are cancelled, citations and references are matched)",
        "btn_start": "START",
        "btn_cancel": "CANCEL",
        "msg_missing_word": "Please select a Word file.\n(If extracting APA only, you can select PDF Folder without a Word file)",
        "msg_missing_word_cross": "Please select a Word file for Cross-Check.",
        "msg_missing_pdf": "Please select a PDF folder or check Link Collection.",
        "msg_no_internet": "Internet connection is required for the selected operations (APA / Link Collection).",
        "msg_cancel_q": "Do you want to stop the operation?",
        "msg_success": "Operation Completed!",
        "help_title": "ATÖxGem User Guide",
        "help_text": "ATÖxGem v7.0 (Absolut) 🎓\n\nThis program performs tedious tasks in seconds during your thesis and article writing process.\n\n==================================================\nMODE 1: CROSS-CHECK MODE (If Red Box is Checked)\n==================================================\nGoal: To find out if the in-text citations (APA or Numbered [1, 2] format) are actually in the REFERENCES list. The program auto-detects the style! (Works Offline).\n\nOutput Files:\n1_Cift_Yonlu_Kontrol_Edilmis_Metin.docx:\n   🟩 Green: Citation found perfectly in references.\n   🟨 Yellow: There is a slight typo, missing number, or a 1-2 year deviation.\n   🟥 Red: Citations with forgotten references or Phantom References.\n\nCapraz_Kontrol_Hata_Raporu.txt:\n   Detailed list of errors.\n\n==================================================\nMODE 2: ARCHIVE SCAN & INTERNET LINK COLLECTION \n==================================================\nFeature A) PDF Archive Matching (If Word + Folder Selected):\nReads citations and searches for them in your archive. Copies found PDFs to the \"1_Bulunan_PDFler\" folder. It also creates a colored Word file (1_Gorsel_Atif_ve_Kaynakca_Analizi.docx) for you to check the matching:\n   🟩 Green (In-text): PDF of the citation was found in the archive.\n   🟥 Red (In-text): PDF of the citation was NOT FOUND in the archive.\n\nFeature B) APA Reference Extractor:\nReads the content of PDFs fetched from the folder, verifies via internet, and creates an APA 7 Word file. (Requires Internet).\nColor Codes (For APA Word File):\n   🟩 Green: Source read from DOI inside PDF, 100% matched.\n   🟨 Yellow: Source found by searching article/author name.\n   🟥 Red: Source NOT FOUND on the internet.\n\nFeature C) Download Link Collection:\nReads REFERENCES list and searches internet for download links. Creates an HTML panel to download them all with one click.\nColor Codes (For Main Word File References):\n   🟩 Green: Source found directly via \"DOI\".\n   🟨 Yellow: No DOI, found via \"Title\".\n   🟥 Red: Article not found.\n\n==================================================\nMODE 3: STANDALONE APA MODE \n==================================================\nIf you select only \"PDF Archive Folder\" and check \"Extract APA References\" without a Word file, it extracts APA for all PDFs in that folder.\n\nWe wish you success in your studies. 🚀"
    },
    "ES": {
        "word_lbl": "1. Archivo Word:",
        "pdf_lbl": "2. Carpeta de Archivo PDF:",
        "btn_select": "Elegir",
        "chk_apa": "Extraer Referencias APA de los PDF Encontrados",
        "chk_indirme": "Escanear REFERENCIAS y Recopilar Enlaces de Descarga",
        "lbl_custom": "Enlace de Origen (ej: https://.../):",
        "chk_kontrol": "Solo Realizar Verificación Cruzada Citas/Referencias",
        "lbl_kontrol_desc": "(Al seleccionar este modo, se cancelan las operaciones anteriores, solo cruza datos)",
        "btn_start": "INICIAR",
        "btn_cancel": "CANCELAR",
        "msg_missing_word": "Por favor seleccione un archivo Word.\n(Si solo extrae APA, puede elegir la carpeta PDF sin un archivo Word)",
        "msg_missing_word_cross": "Seleccione un archivo Word para la Verificación Cruzada.",
        "msg_missing_pdf": "Seleccione una carpeta PDF o marque Recopilación de Enlaces.",
        "msg_no_internet": "Se requiere conexión a Internet para las operaciones seleccionadas (APA / Enlaces).",
        "msg_cancel_q": "¿Desea detener la operación?",
        "msg_success": "¡Operación Completada!",
        "help_title": "Guía de Usuario ATÖxGem",
        "help_text": "ATÖxGem v7.0 (Absolut) 🎓\n\nEste programa realiza tareas tediosas en segundos durante su proceso de escritura.\n\n==================================================\nMODO 1: VERIFICACIÓN CRUZADA (Caja Roja Seleccionada)\n==================================================\nObjetivo: Comprobar si las citas (estilo APA o Numerado [1, 2]) están realmente en la lista de REFERENCIAS. ¡Autodetecta el estilo! (Funciona sin conexión).\n\n¿Cuáles son los archivos de salida?\n1_Cift_Yonlu_Kontrol_Edilmis_Metin.docx:\n   🟩 Verde: Cita encontrada perfectamente en las referencias.\n   🟨 Amarillo: Hay un ligero error tipográfico, número faltante o desviación de años.\n   🟥 Rojo: Citas con referencias olvidadas o Referencias Fantasma.\n\nCapraz_Kontrol_Hata_Raporu.txt:\n   Lista detallada de errores.\n\n==================================================\nMODO 2: ESCANEO DE ARCHIVO Y RECOPILACIÓN DE ENLACES\n==================================================\nCaracterística A) Búsqueda en Archivo PDF (Si elige Word + Carpeta):\nLee citas y las busca en su archivo. Copia los PDF encontrados a la carpeta \"1_Bulunan_PDFler\". También crea un archivo Word coloreado (1_Gorsel_Atif_ve_Kaynakca_Analizi.docx):\n   🟩 Verde (En texto): PDF de la cita encontrado en el archivo.\n   🟥 Rojo (En texto): PDF de la cita NO ENCONTRADO en el archivo.\n\nCaracterística B) Extractor de Referencias APA:\nLee el contenido de los PDF, verifica por internet y crea un archivo Word APA 7. (Requiere Internet).\nCódigos de Color (Para archivo Word APA):\n   🟩 Verde: Origen leído del DOI en el PDF, 100% coincidencia.\n   🟨 Amarillo: Origen encontrado buscando nombre de artículo/autor.\n   🟥 Rojo: Origen NO ENCONTRADO en internet.\n\nCaracterística C) Recopilación de Enlaces de Descarga:\nLee la lista de REFERENCIAS y busca enlaces de descarga. Crea un panel HTML para descargar todo.\nCódigos de Color (Para Referencias del Word Principal):\n   🟩 Verde: Origen encontrado directamente vía \"DOI\".\n   🟨 Amarillo: Sin DOI, encontrado vía \"Título\".\n   🟥 Rojo: Artículo no encontrado.\n\n==================================================\nMODO 3: MODO APA INDEPENDIENTE\n==================================================\nSi selecciona solo \"Carpeta de Archivo PDF\" y marca \"Extraer Referencias APA\", extrae el APA para todos los PDF en esa carpeta.\n\nLe deseamos éxito en sus estudios. 🚀"
    },
    "FR": {
        "word_lbl": "1. Fichier Word :",
        "pdf_lbl": "2. Dossier d'Archives PDF :",
        "btn_select": "Choisir",
        "chk_apa": "Extraire les Références APA des PDF Trouvés",
        "chk_indirme": "Scanner les RÉFÉRENCES et Collecter les Liens",
        "lbl_custom": "Lien Source (ex: https://.../) :",
        "chk_kontrol": "Uniquement Vérification Croisée Citations/Références",
        "lbl_kontrol_desc": "(Ce mode annule les opérations ci-dessus, croise uniquement le texte et les références)",
        "btn_start": "DÉMARRER",
        "btn_cancel": "ANNULER",
        "msg_missing_word": "Veuillez sélectionner un fichier Word.\n(Pour APA seul, sélectionnez le dossier PDF sans Word)",
        "msg_missing_word_cross": "Veuillez sélectionner un fichier Word pour la Vérification.",
        "msg_missing_pdf": "Sélectionnez un dossier PDF ou cochez Collecte de Liens.",
        "msg_no_internet": "Connexion Internet requise pour ces opérations (APA / Liens).",
        "msg_cancel_q": "Voulez-vous arrêter l'opération ?",
        "msg_success": "Opération Terminée !",
        "help_title": "Guide de l'Utilisateur ATÖxGem",
        "help_text": "ATÖxGem v7.0 (Absolut) 🎓\n\nCe programme effectue des tâches fastidieuses en quelques secondes lors de la rédaction.\n\n==================================================\nMODE 1 : VÉRIFICATION CROISÉE (Case Rouge Cochée)\n==================================================\nObjectif : Vérifier si les citations (style APA ou Numéroté [1, 2]) sont réellement dans les RÉFÉRENCES. Détection automatique du style ! (Hors ligne).\n\nFichiers de sortie :\n1_Cift_Yonlu_Kontrol_Edilmis_Metin.docx :\n   🟩 Vert : Citation trouvée parfaitement dans les références.\n   🟨 Jaune : Légère faute de frappe ou numéro manquant.\n   🟥 Rouge : Citations sans références ou Références Fantômes.\n\nCapraz_Kontrol_Hata_Raporu.txt :\n   Liste détaillée des erreurs.\n\n==================================================\nMODE 2 : SCAN D'ARCHIVE & COLLECTE DE LIENS\n==================================================\nFonctionnalité A) Correspondance d'Archives PDF (Word + Dossier) :\nLit les citations et les cherche dans l'archive. Copie les PDF trouvés. Crée un fichier Word coloré (1_Gorsel_Atif_ve_Kaynakca_Analizi.docx) :\n   🟩 Vert (Dans le texte) : PDF de la citation trouvé dans l'archive.\n   🟥 Rouge (Dans le texte) : PDF de la citation INTROUVABLE.\n\nFonctionnalité B) Extracteur de Références APA :\nLit le contenu des PDF, vérifie sur internet et crée un fichier Word APA 7. (Internet requis).\nCodes Couleurs (Pour Word APA) :\n   🟩 Vert : Source lue depuis le DOI, 100% de correspondance.\n   🟨 Jaune : Source trouvée en cherchant le nom.\n   🟥 Rouge : Source INTROUVABLE sur internet.\n\nFonctionnalité C) Collecte de Liens de Téléchargement :\nCherche les liens de téléchargement sur internet. Crée un panneau HTML.\nCodes Couleurs (Pour Références Word Principal) :\n   🟩 Vert : Trouvée via \"DOI\".\n   🟨 Jaune : Trouvée via \"Titre\".\n   🟥 Rouge : Introuvable.\n\n==================================================\nMODE 3 : MODE APA AUTONOME\n==================================================\nSélectionnez uniquement \"Dossier d'Archives PDF\" et cochez \"Extraire les Références APA\" pour extraire l'APA des PDF du dossier.\n\nNous vous souhaitons beaucoup de succès. 🚀"
    },
    "DE": {
        "word_lbl": "1. Word-Datei:",
        "pdf_lbl": "2. PDF-Archivordner:",
        "btn_select": "Wählen",
        "chk_apa": "APA-Referenzen aus gefundenen PDFs extrahieren",
        "chk_indirme": "REFERENZEN scannen und Download-Links sammeln",
        "lbl_custom": "Benutzerdefinierter Quell-Link (z.B. https://.../):",
        "chk_kontrol": "Nur Kreuzprüfung von Zitaten / Referenzen",
        "lbl_kontrol_desc": "(Dieser Modus storniert obige Vorgänge, vergleicht nur Text und Referenzen)",
        "btn_start": "STARTEN",
        "btn_cancel": "ABBRECHEN",
        "msg_missing_word": "Bitte wählen Sie eine Word-Datei aus.\n(Für reines APA wählen Sie nur den PDF-Ordner)",
        "msg_missing_word_cross": "Bitte wählen Sie eine Word-Datei für die Kreuzprüfung.",
        "msg_missing_pdf": "Wählen Sie einen PDF-Ordner oder kreuzen Sie Link-Sammlung an.",
        "msg_no_internet": "Für diese Vorgänge (APA / Links) ist Internet erforderlich.",
        "msg_cancel_q": "Möchten Sie den Vorgang abbrechen?",
        "msg_success": "Vorgang Abgeschlossen!",
        "help_title": "ATÖxGem Benutzerhandbuch",
        "help_text": "ATÖxGem v7.0 (Absolut) 🎓\n\nDieses Programm erledigt lästige Aufgaben während Ihres Schreibprozesses in Sekundenschnelle.\n\n==================================================\nMODUS 1: KREUZPRÜFUNG (Rotes Kästchen Markiert)\n==================================================\nZiel: Herauszufinden, ob Textzitate (APA oder Nummeriert [1, 2]) tatsächlich im LITERATURVERZEICHNIS stehen. Stil wird automatisch erkannt! (Offline).\n\nAusgabedateien:\n1_Cift_Yonlu_Kontrol_Edilmis_Metin.docx:\n   🟩 Grün: Zitat perfekt in den Referenzen gefunden.\n   🟨 Gelb: Leichter Tippfehler oder fehlende Nummer.\n   🟥 Rot: Zitate mit vergessenen Referenzen oder Geisterreferenzen.\n\nCapraz_Kontrol_Hata_Raporu.txt:\n   Detaillierte Fehlerliste.\n\n==================================================\nMODUS 2: ARCHIV-SCAN & LINK-SAMMLUNG\n==================================================\nFunktion A) PDF-Archivabgleich (Wenn Word + Ordner):\nSucht Zitate im Archiv und kopiert PDFs. Erstellt eine farbige Word-Datei (1_Gorsel_Atif_ve_Kaynakca_Analizi.docx):\n   🟩 Grün (Im Text): PDF des Zitats im Archiv gefunden.\n   🟥 Rot (Im Text): PDF des Zitats im Archiv NICHT GEFUNDEN.\n\nFunktion B) APA-Referenz-Extraktor:\nLiest PDFs, verifiziert via Internet und erstellt eine APA 7 Word-Datei. (Internet erforderlich).\nFarbcodes (Für APA Word-Datei):\n   🟩 Grün: 100% Übereinstimmung über DOI.\n   🟨 Gelb: Über Artikel-/Autorennamen gefunden.\n   🟥 Rot: Quelle NICHT GEFUNDEN.\n\nFunktion C) Download-Link-Sammlung:\nSucht im Internet nach Links. Erstellt ein HTML-Panel.\nFarbcodes (Für Haupt-Word-Datei):\n   🟩 Grün: Direkt über \"DOI\" gefunden.\n   🟨 Gelb: Über \"Titel\" gefunden.\n   🟥 Rot: Artikel nicht gefunden.\n\n==================================================\nMODUS 3: STANDALONE APA MODUS\n==================================================\nWählen Sie nur \"PDF-Archivordner\" aus und kreuzen Sie \"APA-Referenzen extrahieren\" an.\n\nWir wünschen Ihnen viel Erfolg. 🚀"
    }
}

# =============================================================================
#                           ARAYÜZ (GUI) VE İZOLASYON MANTIĞI
# =============================================================================

islem_calisiyor = False

def ui_mod_guncelle():
    if kontrol_var.get():
        ent_klasor.config(state=tk.DISABLED)
        btn_klasor.config(state=tk.DISABLED)
        chk_indirme.config(state=tk.DISABLED)
        chk_apa.config(state=tk.DISABLED) 
        ent_custom.config(state=tk.DISABLED)
        indirme_var.set(False)
        apa_var.set(False)
    else:
        ent_klasor.config(state=tk.NORMAL)
        btn_klasor.config(state=tk.NORMAL)
        chk_indirme.config(state=tk.NORMAL)
        chk_apa.config(state=tk.NORMAL)
        ent_custom.config(state=tk.NORMAL)
    kontrol_et_ve_guncelle()

def toggle_islem():
    global islem_calisiyor
    lang = lang_var.get().split(" ")[1]
    t = TRANSLATIONS[lang]
    
    if not islem_calisiyor:
        w = dosya_var.get()
        p = klasor_var.get()
        indirme = indirme_var.get()
        apa = apa_var.get()
        c_url = custom_url_var.get()
        k_aktif = kontrol_var.get()
        
        internet_gerekli = (indirme or apa)
        is_standalone_apa = (not k_aktif) and apa and p and not w and not indirme

        if internet_gerekli and not internet_var_mi():
             messagebox.showerror("Hata/Error", t["msg_no_internet"])
             return
             
        if k_aktif:
            if not w:
                messagebox.showwarning("Eksik/Missing", t["msg_missing_word_cross"])
                return
        else:
            if is_standalone_apa:
                pass 
            elif not w:
                messagebox.showwarning("Eksik/Missing", t["msg_missing_word"])
                return
            elif not p and not indirme:
                messagebox.showwarning("Eksik/Missing", t["msg_missing_pdf"])
                return

        islem_calisiyor = True
        btn_baslat.config(text=t["btn_cancel"], bg=GOOGLE_RED)
        threading.Thread(target=tam_otomatik_islem, args=(w, p, indirme, apa, c_url, k_aktif), daemon=True).start()
    else:
        if messagebox.askyesno("Soru/Question", t["msg_cancel_q"]):
            stop_event.set()
            log_yaz("İptal ediliyor... / Cancelling...")

def reset_ui():
    global islem_calisiyor
    islem_calisiyor = False
    stop_event.clear()
    lang = lang_var.get().split(" ")[1]
    btn_baslat.config(text=TRANSLATIONS[lang]["btn_start"], bg=GOOGLE_GREEN)
    kontrol_et_ve_guncelle()

def kontrol_et_ve_guncelle():
    durum = internet_var_mi()
    internet_gerekli = (indirme_var.get() or apa_var.get())
    k_aktif = kontrol_var.get()
    
    lang = lang_var.get().split(" ")[1]
    t = TRANSLATIONS[lang]
    
    if durum:
        lbl_net.config(text="● Online", fg=GOOGLE_GREEN)
        if not islem_calisiyor:
            btn_baslat.config(state=tk.NORMAL)
            if btn_baslat.cget('bg') == "gray":
                btn_baslat.config(bg=GOOGLE_GREEN, text=t["btn_start"])
    else:
        lbl_net.config(text="● Offline", fg=GOOGLE_RED)
        if not islem_calisiyor:
            if not internet_gerekli:
                btn_baslat.config(state=tk.NORMAL)
                if btn_baslat.cget('bg') == "gray":
                    btn_baslat.config(bg=GOOGLE_GREEN, text=t["btn_start"])
            else:
                btn_baslat.config(state=tk.DISABLED, bg="gray", text=t["btn_start"])
    
    global loop_id
    try: root.after_cancel(loop_id)
    except: pass
    loop_id = root.after(3000, kontrol_et_ve_guncelle)

def dosya_sec():
    f = filedialog.askopenfilename(filetypes=[("Word", "*.docx")])
    if f: dosya_var.set(f)

def klasor_sec():
    d = filedialog.askdirectory()
    if d: klasor_var.set(d)

def degistir_dil(secim):
    lang = secim.split(" ")[1]
    t = TRANSLATIONS[lang]
    
    lbl_word.config(text=t["word_lbl"])
    lbl_pdf.config(text=t["pdf_lbl"])
    btn_word.config(text=t["btn_select"])
    btn_pdf.config(text=t["btn_select"])
    chk_apa_btn.config(text=t["chk_apa"])
    chk_indirme_btn.config(text=t["chk_indirme"])
    lbl_custom.config(text=t["lbl_custom"])
    chk_kontrol_btn.config(text=t["chk_kontrol"])
    lbl_kontrol_desc.config(text=t["lbl_kontrol_desc"])
    
    if not islem_calisiyor:
        btn_baslat.config(text=t["btn_start"])
    else:
        btn_baslat.config(text=t["btn_cancel"])

# =============================================================================
# BİLGİ (YARDIM) PENCERESİ FONKSİYONU
# =============================================================================
def yardim_penceresi_ac():
    lang = lang_var.get().split(" ")[1]
    t = TRANSLATIONS[lang]
    
    yardim_win = tk.Toplevel(root)
    yardim_win.title(t["help_title"])
    pencereyi_ortala(yardim_win, 700, 650)
    yardim_win.configure(bg=BG_COLOR)
    
    yardim_text = scrolledtext.ScrolledText(yardim_win, wrap=tk.WORD, font=("Arial", 10), bg="#f8f9fa", fg="#202124", padx=15, pady=15)
    yardim_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
    
    yardim_text.insert(tk.END, t["help_text"])
    yardim_text.config(state=tk.DISABLED)

# =============================================================================
#                           ANA ARAYÜZ (GUI)
# =============================================================================

root = tk.Tk()
root.title("ATÖxGem v7.0 (Absolut)")
root.configure(bg=BG_COLOR)
pencereyi_ortala(root, 650, 680)

try: root.iconbitmap(resource_path("icon.ico"))
except: pass

# --- HEADER FRAME & BİLGİ/DİL BUTONU ---
header_frame = tk.Frame(root, bg=BG_COLOR)
header_frame.pack(fill="x", pady=10)

# SOL ÜST KÖŞE (Dil Seçimi ve Online Durumu)
left_frame = tk.Frame(header_frame, bg=BG_COLOR)
left_frame.pack(side=tk.LEFT, padx=10)

lang_var = tk.StringVar(value="友 TR")
lang_menu = tk.OptionMenu(left_frame, lang_var, "友 TR", "友 EN", "友 ES", "友 FR", "友 DE", command=degistir_dil)
lang_menu.config(bg="#e8f0fe", fg=GOOGLE_BLUE, bd=1, relief=tk.RAISED, font=("Arial", 8, "bold"), cursor="hand2")
lang_menu.pack(anchor="w", pady=(0, 2))

lbl_net = tk.Label(left_frame, text="● ...", font=("Arial", 9, "bold"), bg=BG_COLOR, fg="gray")
lbl_net.pack(anchor="w")

# ORTA KISIM (Başlık)
title_frame = tk.Frame(header_frame, bg=BG_COLOR)
title_frame.pack(side=tk.LEFT, expand=True)

tk.Label(title_frame, text="ATÖ", font=("Arial Black", 24), fg=GOOGLE_BLUE, bg=BG_COLOR).pack(side=tk.LEFT)
tk.Label(title_frame, text="x", font=("Arial Black", 24), fg=GOOGLE_RED, bg=BG_COLOR).pack(side=tk.LEFT)
tk.Label(title_frame, text="Gem", font=("Arial Black", 24), fg=GOOGLE_YELLOW, bg=BG_COLOR).pack(side=tk.LEFT)

# SAĞ ÜST KÖŞE ("i" Butonu)
right_frame = tk.Frame(header_frame, bg=BG_COLOR)
right_frame.pack(side=tk.RIGHT, padx=20)

btn_info = tk.Button(right_frame, text=" i ", font=("Arial", 12, "bold", "italic"), bg="#e8f0fe", fg=GOOGLE_BLUE, bd=1, relief=tk.RAISED, command=yardim_penceresi_ac, cursor="hand2")
btn_info.pack(pady=(0, 5))

# GİRDİ ALANLARI
input_frame = tk.Frame(root, bg=BG_COLOR, pady=10)
input_frame.pack(fill="x", padx=40)

dosya_var = tk.StringVar()
klasor_var = tk.StringVar()

# Başlangıçta tüm kutucuklar kapalı
indirme_var = tk.BooleanVar(value=False) 
apa_var = tk.BooleanVar(value=False) 
custom_url_var = tk.StringVar() 
kontrol_var = tk.BooleanVar(value=False) 

lbl_word = tk.Label(input_frame, text=TRANSLATIONS["TR"]["word_lbl"], bg=BG_COLOR)
lbl_word.pack(anchor="w")
tk.Entry(input_frame, textvariable=dosya_var, bg="#f1f3f4").pack(fill="x", pady=5)
btn_word = tk.Button(input_frame, text=TRANSLATIONS["TR"]["btn_select"], command=dosya_sec, bg=GOOGLE_BLUE, fg="white")
btn_word.pack(anchor="e")

lbl_pdf = tk.Label(input_frame, text=TRANSLATIONS["TR"]["pdf_lbl"], bg=BG_COLOR)
lbl_pdf.pack(anchor="w", pady=(10,0))
ent_klasor = tk.Entry(input_frame, textvariable=klasor_var, bg="#f1f3f4")
ent_klasor.pack(fill="x", pady=5)
btn_pdf = tk.Button(input_frame, text=TRANSLATIONS["TR"]["btn_select"], command=klasor_sec, bg=GOOGLE_RED, fg="white")
btn_pdf.pack(anchor="e")

check_frame = tk.Frame(root, bg=BG_COLOR)
check_frame.pack(fill="x", padx=40, pady=5)

chk_apa_btn = tk.Checkbutton(check_frame, text=TRANSLATIONS["TR"]["chk_apa"], 
               variable=apa_var, command=kontrol_et_ve_guncelle, bg=BG_COLOR, font=("Arial", 9, "bold"))
chk_apa_btn.pack(anchor="w", pady=(0, 5))

chk_indirme_btn = tk.Checkbutton(check_frame, text=TRANSLATIONS["TR"]["chk_indirme"], 
               variable=indirme_var, command=kontrol_et_ve_guncelle, bg=BG_COLOR, font=("Arial", 9, "bold"))
chk_indirme_btn.pack(anchor="w")

custom_url_frame = tk.Frame(root, bg=BG_COLOR)
custom_url_frame.pack(fill="x", padx=60, pady=2) 
lbl_custom = tk.Label(custom_url_frame, text=TRANSLATIONS["TR"]["lbl_custom"], bg=BG_COLOR, font=("Arial", 8))
lbl_custom.pack(side=tk.LEFT)
ent_custom = tk.Entry(custom_url_frame, textvariable=custom_url_var, bg="#f1f3f4", width=35)
ent_custom.pack(side=tk.LEFT, padx=5)

ayrac = tk.Frame(root, height=2, bd=1, relief=tk.SUNKEN)
ayrac.pack(fill="x", padx=40, pady=10)

kontrol_frame = tk.Frame(root, bg=BG_COLOR)
kontrol_frame.pack(fill="x", padx=40)
chk_kontrol_btn = tk.Checkbutton(kontrol_frame, text=TRANSLATIONS["TR"]["chk_kontrol"], 
               variable=kontrol_var, command=ui_mod_guncelle, bg=BG_COLOR, font=("Arial", 10, "bold"), fg=GOOGLE_RED)
chk_kontrol_btn.pack(anchor="w")
lbl_kontrol_desc = tk.Label(kontrol_frame, text=TRANSLATIONS["TR"]["lbl_kontrol_desc"], bg=BG_COLOR, font=("Arial", 8, "italic"), fg="gray")
lbl_kontrol_desc.pack(anchor="w", padx=20)

btn_baslat = tk.Button(root, text=TRANSLATIONS["TR"]["btn_start"], command=toggle_islem, bg=GOOGLE_GREEN, fg="white", font=("Arial", 12, "bold"), height=2, width=30)
btn_baslat.pack(pady=10)

log_frame = tk.Frame(root, bg=BG_COLOR, padx=20, pady=10)
log_frame.pack(fill="both", expand=True)
text_log = scrolledtext.ScrolledText(log_frame, height=8, bg="black", fg="#00ff00", state=tk.DISABLED)
text_log.pack(fill="both", expand=True)

loop_id = None
degistir_dil("友 TR")
kontrol_et_ve_guncelle()
root.mainloop()